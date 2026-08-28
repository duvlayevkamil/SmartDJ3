"""
Chop etish — sozlamalar + HTML preview + QPrintPreviewDialog
"""
from PyQt6.QtWidgets import (QDialog, QVBoxLayout, QHBoxLayout, QPushButton,
                             QLabel, QComboBox, QLineEdit, QGroupBox,
                             QGridLayout, QRadioButton, QButtonGroup,
                             QSlider, QWidget, QTextEdit, QSplitter,
                             QMessageBox, QFileDialog)
from PyQt6.QtPrintSupport import QPrintPreviewDialog, QPrinter
from PyQt6.QtGui import QPageLayout, QPageSize, QTextDocument
from PyQt6.QtCore import Qt, QMarginsF
from datetime import datetime
import logging

PERIODS_PER_DAY = 6  # Kuniga maksimal 6 dars (7-dars yo'q)

from core.exporter import build_html

KUNLAR = ["Dushanba", "Seshanba", "Chorshanba", "Payshanba", "Juma", "Shanba"]


class ExportSettingsDialog(QDialog):
    def __init__(self, db, tt, classes, parent=None, export_format='pdf', tt2=None):
        super().__init__(parent)
        self.db = db
        self.tt = tt
        self.tt2 = tt2 if tt2 else {}
        self.classes = classes
        self.fmt = export_format

        fmt_labels = {'pdf': 'PDF', 'excel': 'Excel', 'word': 'Word', 'html': 'HTML', 'csv': 'CSV'}
        fmt_icons = {'pdf': '📄', 'excel': '📊', 'word': '📝', 'html': '🌐', 'csv': '📋'}
        label = fmt_labels.get(export_format, export_format.upper())
        icon = fmt_icons.get(export_format, '📄')
        self.setWindowTitle(f"Chop etish — {label}")
        self.setWindowFlags(Qt.WindowType.Window | Qt.WindowType.WindowMinimizeButtonHint
                           | Qt.WindowType.WindowMaximizeButtonHint | Qt.WindowType.WindowCloseButtonHint)
        self.setMinimumSize(1000, 700)
        self.resize(1200, 850)

        self._build()
        self._refresh()

    def _build(self):
        root = QVBoxLayout()
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)
        self.setLayout(root)

        # Header
        hdr = QWidget()
        hdr.setFixedHeight(40)
        hdr.setStyleSheet("background:#2C3E50;")
        hl = QHBoxLayout()
        hl.setContentsMargins(15, 0, 15, 0)
        hdr.setLayout(hl)
        fmt_icons = {'pdf': '📄', 'excel': '📊', 'word': '📝', 'html': '🌐', 'csv': '📋'}
        fmt_labels = {'pdf': 'PDF', 'excel': 'Excel', 'word': 'Word', 'html': 'HTML', 'csv': 'CSV'}
        ico = fmt_icons.get(self.fmt, '📄')
        label = fmt_labels.get(self.fmt, self.fmt.upper())
        lbl = QLabel(f"{ico} Chop etish — {label}")
        lbl.setStyleSheet("color:white; font-size:15px; font-weight:bold;")
        hl.addWidget(lbl)
        hl.addStretch()
        info = QLabel(f"{len(self.classes)} sinf | {len(self.tt)} dars")
        info.setStyleSheet("color:#BDC3C7; font-size:12px;")
        hl.addWidget(info)
        root.addWidget(hdr)

        # Splitter
        sp = QSplitter(Qt.Orientation.Horizontal)
        sp.setHandleWidth(2)

        # Left: sozlamalar
        left = QWidget()
        left.setFixedWidth(280)
        left.setStyleSheet("background:#F8F9FA; border-right:1px solid #DEE2E6;")
        lv = QVBoxLayout()
        lv.setContentsMargins(8, 8, 8, 8)
        lv.setSpacing(5)
        left.setLayout(lv)

        lv.addWidget(self._hdr("Sozlamalar"))

        # Chop etish turi
        g = self._grp("Chop etish turi")
        gl = QVBoxLayout()
        g.setLayout(gl)
        self.type_cb = QComboBox()
        self.type_cb.addItems(["Umumiy jadval", "Sinf bo'yicha",
                               "O'qituvchi bo'yicha", "Fan bo'yicha"])
        self.type_cb.setStyleSheet(
            "QComboBox{padding:4px;font-size:11px;border:1px solid #BDC3C7;border-radius:3px;}"
            "QComboBox:hover{background:#3498DB;color:#000000;font-weight:bold;}"
            "QComboBox::drop-down{border:none;}"
            "QComboBox QAbstractItemView{background:white;selection-background-color:#3498DB;selection-color:#000000;font-weight:bold;}")
        self.type_cb.currentIndexChanged.connect(self._type_chg)
        gl.addWidget(self.type_cb)
        self.item_cb = QComboBox()
        self.item_cb.setStyleSheet("padding:4px;font-size:11px;border:1px solid #BDC3C7;border-radius:3px;")
        self.item_cb.currentIndexChanged.connect(self._refresh)
        self.item_cb.setVisible(False)
        gl.addWidget(self.item_cb)
        lv.addWidget(g)

        # Hafta tanlash
        g = self._grp("Hafta tanlash")
        gl = QVBoxLayout()
        g.setLayout(gl)
        self.week_cb = QComboBox()
        self.week_cb.addItems(["1-hafta (Toq — Numerator)", "2-hafta (Juft — Denominator)", "Ikkalasi (2-haftalik)"])
        self.week_cb.setStyleSheet(
            "QComboBox{padding:4px;font-size:11px;border:1px solid #BDC3C7;border-radius:3px;}"
            "QComboBox:hover{background:#3498DB;color:#000000;font-weight:bold;}"
            "QComboBox::drop-down{border:none;}"
            "QComboBox QAbstractItemView{background:white;selection-background-color:#3498DB;selection-color:#000000;font-weight:bold;}")
        self.week_cb.currentIndexChanged.connect(self._refresh)
        gl.addWidget(self.week_cb)
        # Agar 2-hafta ma'lumoti bo'lmasa, "Ikkalasi" ni o'chirish
        if not self.tt2:
            self.week_cb.setItemEnabled(1, False)
            self.week_cb.setItemEnabled(2, False)
        lv.addWidget(g)

        # Maktab nomi
        g = self._grp("Maktab nomi")
        gl = QVBoxLayout()
        g.setLayout(gl)
        self.school = QLineEdit("Umumiy o'rta ta'lim maktabi")
        self.school.setStyleSheet("padding:4px;font-size:11px;border:1px solid #BDC3C7;border-radius:3px;")
        self.school.textChanged.connect(self._refresh)
        gl.addWidget(self.school)
        lv.addWidget(g)

        # Shrift o'lchami
        g = self._grp("Shrift o'lchami")
        gl = QHBoxLayout()
        g.setLayout(gl)
        self.font_sl = QSlider(Qt.Orientation.Horizontal)
        self.font_sl.setMinimum(6)
        self.font_sl.setMaximum(14)
        self.font_sl.setValue(9)
        self.font_sl.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.font_sl.setTickInterval(1)
        self.font_sl.valueChanged.connect(self._font_chg)
        gl.addWidget(self.font_sl)
        self.font_lbl = QLabel("9pt")
        self.font_lbl.setStyleSheet("font-weight:bold;color:#2C3E50;min-width:35px;")
        self.font_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        gl.addWidget(self.font_lbl)
        lv.addWidget(g)

        # Qog'oz
        g = self._grp("Qog'oz formati")
        gl = QGridLayout()
        g.setLayout(gl)
        gl.addWidget(QLabel("Format:"), 0, 0)
        self.paper_cb = QComboBox()
        self.paper_cb.addItems(["A4", "A3", "A2", "Letter", "Legal"])
        self.paper_cb.currentTextChanged.connect(self._refresh)
        gl.addWidget(self.paper_cb, 0, 1)
        self.portrait = QRadioButton("Vertikal")
        self.landscape = QRadioButton("Gorizontal")
        self.landscape.setChecked(True)
        bg = QButtonGroup()
        bg.addButton(self.portrait)
        bg.addButton(self.landscape)
        self.portrait.toggled.connect(self._refresh)
        gl.addWidget(self.portrait, 1, 0)
        gl.addWidget(self.landscape, 1, 1)
        lv.addWidget(g)

        # Chegaralar
        g = self._grp("Chegaralar (mm)")
        gl = QGridLayout()
        g.setLayout(gl)
        self.mt = QLineEdit("10"); self.mb = QLineEdit("10")
        self.ml = QLineEdit("10"); self.mr = QLineEdit("10")
        for e in [self.mt, self.mb, self.ml, self.mr]:
            e.setAlignment(Qt.AlignmentFlag.AlignCenter)
            e.setFixedWidth(45)
            e.editingFinished.connect(self._refresh)
        gl.addWidget(QLabel("Yuqori:"), 0, 0); gl.addWidget(self.mt, 0, 1)
        gl.addWidget(QLabel("Past:"), 0, 2); gl.addWidget(self.mb, 0, 3)
        gl.addWidget(QLabel("Chap:"), 1, 0); gl.addWidget(self.ml, 1, 1)
        gl.addWidget(QLabel("O'ng:"), 1, 2); gl.addWidget(self.mr, 1, 3)
        lv.addWidget(g)

        lv.addStretch()

        # Tugmalar
        bl = QVBoxLayout()
        bl.setSpacing(6)

        preview_btn = QPushButton("👁️ To'liq ko'rish")
        preview_btn.setStyleSheet(
            "QPushButton{background:#3498DB;color:white;padding:10px;font-size:13px;"
            "border-radius:5px;font-weight:bold;border:none;}"
            "QPushButton:hover{background:#2980B9;}")
        preview_btn.clicked.connect(self._full_preview)
        bl.addWidget(preview_btn)

        fmt_labels = {'pdf': 'PDF', 'excel': 'Excel', 'word': 'Word', 'html': 'HTML', 'csv': 'CSV'}
        save_label = fmt_labels.get(self.fmt, self.fmt.upper())
        save_btn = QPushButton(f"💾 {save_label} saqlash")
        save_btn.setStyleSheet(
            "QPushButton{background:#27AE60;color:white;padding:10px;font-size:13px;"
            "border-radius:5px;font-weight:bold;border:none;}"
            "QPushButton:hover{background:#229954;}")
        save_btn.clicked.connect(self._save)
        bl.addWidget(save_btn)

        cancel_btn = QPushButton("Bekor")
        cancel_btn.setStyleSheet(
            "QPushButton{background:#95A5A6;color:white;padding:10px;font-size:13px;"
            "border-radius:5px;font-weight:bold;border:none;}"
            "QPushButton:hover{background:#7F8C8D;}")
        cancel_btn.clicked.connect(self.reject)
        bl.addWidget(cancel_btn)

        lv.addLayout(bl)

        sp.addWidget(left)

        # Right: HTML preview
        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        self.preview.setStyleSheet("background:#e0e0e0; border:none; color:#000000;")
        sp.addWidget(self.preview)

        sp.setSizes([280, 800])
        root.addWidget(sp, 1)

    def _hdr(self, t):
        l = QLabel(t)
        l.setStyleSheet("font-size:14px;font-weight:bold;color:#2C3E50;padding:4px;")
        return l

    def _grp(self, t):
        g = QGroupBox(t)
        g.setStyleSheet("QGroupBox{font-size:11px;font-weight:bold;border:1px solid #DEE2E6;"
                        "border-radius:4px;margin-top:8px;padding-top:14px;}")
        return g

    # --- Signals ---
    def _type_chg(self, i):
        self.item_cb.blockSignals(True)
        self.item_cb.clear()
        if i == 0:
            self.item_cb.setVisible(False)
        elif i == 1:
            self.item_cb.setVisible(True)
            for c in self.classes:
                self.item_cb.addItem(c[1], c[0])
        elif i == 2:
            self.item_cb.setVisible(True)
            for t in self.db.get_all_teachers():
                self.item_cb.addItem(t[1], t[0])
        elif i == 3:
            self.item_cb.setVisible(True)
            for s in self.db.get_all_subjects():
                self.item_cb.addItem(s[1], s[0])
        self.item_cb.blockSignals(False)
        self._refresh()

    def _font_chg(self, v):
        self.font_lbl.setText(f"{v}pt")
        self._refresh()

    def _refresh(self):
        html = self._build_html()
        self.preview.setHtml(html)

    # --- Helpers ---
    def _get_type(self):
        i = self.type_cb.currentIndex()
        if i == 0: return 'umumiy', None, ''
        if i == 1: return 'sinf', self.item_cb.currentData(), self.item_cb.currentText()
        if i == 2: return 'ustoz', self.item_cb.currentData(), self.item_cb.currentText()
        if i == 3: return 'fan', self.item_cb.currentText(), self.item_cb.currentText()
        return 'umumiy', None, ''

    def _get_active_tt(self):
        """Tanlangan haftaning ma'lumotini qaytaradi (preview uchun)"""
        week_idx = self.week_cb.currentIndex()
        if week_idx == 0:
            return self.tt
        elif week_idx == 1 and self.tt2:
            return self.tt2
        elif week_idx == 2:
            # Ikkalasi — 1-haftani qaytarish (preview uchun)
            return self.tt
        return self.tt

    def _get_export_tt(self):
        """Export uchun to'g'ri ma'lumotni qaytaradi"""
        week_idx = self.week_cb.currentIndex()
        if week_idx == 0:
            return [self.tt]
        elif week_idx == 1 and self.tt2:
            return [self.tt2]
        elif week_idx == 2:
            return [self.tt, self.tt2]
        return [self.tt]

    def _build_html(self):
        et, eid, ename = self._get_type()
        week_idx = self.week_cb.currentIndex()
        
        # Debug: tt2 hajmini tekshirish
        tt2_len = len(self.tt2) if self.tt2 else 0
        
        if week_idx == 2:
            # Ikkalasi — ikki alohida jadval
            html1 = build_html(
                self.tt, self.classes,
                etype=et, eid=eid, ename=ename,
                school=self.school.text(),
                fs=self.font_sl.value()
            )
            if tt2_len > 0:
                html2 = build_html(
                    self.tt2, self.classes,
                    etype=et, eid=eid, ename=ename,
                    school=self.school.text(),
                    fs=self.font_sl.value()
                )
                # Ikki jadvalni birlashtirish
                return html1.replace('</body>', 
                    f'<div style="page-break-before:always;"></div>'
                    f'<h2 style="text-align:center; color:#2C3E50; margin:20px 0;">2-HAFTA (Juft — Denominator)</h2>'
                    + html2.split('<body>')[1].split('</body>')[0]
                    + '</body>')
            else:
                # tt2 bo'sh — faqat 1-hafta
                return html1
        else:
            tt = self._get_active_tt()
            return build_html(
                tt, self.classes,
                etype=et, eid=eid, ename=ename,
                school=self.school.text(),
                fs=self.font_sl.value()
            )

    def _make_printer(self):
        pm = {'A3': QPageSize.PageSizeId.A3, 'A4': QPageSize.PageSizeId.A4,
              'A2': QPageSize.PageSizeId.A2, 'Letter': QPageSize.PageSizeId.Letter,
              'Legal': QPageSize.PageSizeId.Legal}
        printer = QPrinter(QPrinter.PrinterMode.ScreenResolution)
        printer.setPageSize(QPageSize(pm.get(self.paper_cb.currentText(), QPageSize.PageSizeId.A4)))
        orient = (QPageLayout.Orientation.Landscape if self.landscape.isChecked()
                  else QPageLayout.Orientation.Portrait)
        printer.setPageOrientation(orient)
        printer.setPageMargins(QMarginsF(
            float(self.ml.text() or 10), float(self.mt.text() or 10),
            float(self.mr.text() or 10), float(self.mb.text() or 10)
        ))
        return printer

    # --- Full Preview ---
    def _full_preview(self):
        try:
            html = self._build_html()
            printer = self._make_printer()
            preview = QPrintPreviewDialog(printer, self)
            preview.setWindowTitle("Jadvalni to'liq ko'rish")
            preview.paintRequested.connect(lambda p: self._render(p, html))
            preview.exec()
        except Exception as e:
            logging.error(f"Preview: {e}")
            QMessageBox.critical(self, "Xatolik", str(e))

    def _render(self, printer, html):
        doc = QTextDocument()
        doc.setHtml(html)
        doc.print(printer)

    # --- Save ---
    def _save(self):
        ext_map = {'pdf': ('pdf', 'PDF (*.pdf)'), 'excel': ('xlsx', 'Excel (*.xlsx)'),
                   'word': ('docx', 'Word (*.docx)'), 'html': ('html', 'HTML (*.html)')}
        ext, filter_str = ext_map.get(self.fmt, ('pdf', 'PDF (*.pdf)'))
        fn, _ = QFileDialog.getSaveFileName(
            self, f"{self.fmt.upper()} saqlash",
            f"Jadval_{datetime.now().strftime('%Y%m%d')}.{ext}",
            filter_str)
        if not fn:
            return
        try:
            if self.fmt == 'pdf':
                self._save_pdf(fn)
            elif self.fmt == 'excel':
                self._save_excel(fn)
            elif self.fmt == 'word':
                self._save_word(fn)
            elif self.fmt == 'html':
                self._save_html(fn)
            QMessageBox.information(self, "OK", f"{self.fmt.upper()} saqlandi!\n{fn}")
            self.accept()
        except Exception as e:
            logging.error(f"Save: {e}")
            QMessageBox.critical(self, "Xatolik", str(e))

    def _save_pdf(self, fn):
        html = self._build_html()
        printer = self._make_printer()
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(fn)
        doc = QTextDocument()
        doc.setHtml(html)
        doc.print(printer)

    def _save_excel(self, fn):
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        tt_list = self._get_export_tt()
        wb = Workbook()
        wb.remove(wb.active)

        week_names = ["1-hafta (Toq)", "2-hafta (Juft)"]
        for wi, tt_data in enumerate(tt_list):
            ws_name = week_names[wi] if wi < len(week_names) else f"Hafta {wi+1}"
            ws = wb.create_sheet(ws_name)
            hf = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
            hfont = Font(bold=True, color="FFFFFF", size=9)
            pf = PatternFill(start_color="34495E", end_color="34495E", fill_type="solid")
            pfont = Font(bold=True, color="FFFFFF", size=8)
            tb = Border(left=Side(style='thin', color='BDC3C7'),
                        right=Side(style='thin', color='BDC3C7'),
                        top=Side(style='thin', color='BDC3C7'),
                        bottom=Side(style='thin', color='BDC3C7'))

            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=43)
            ws.cell(1, 1, f"DARS JADVALI — {ws_name}").font = Font(bold=True, size=14, color="2C3E50")
            ws.cell(1, 1).alignment = Alignment(horizontal='center', vertical='center')

            ws.cell(3, 1, "Sinf").font = hfont; ws.cell(3, 1).fill = hf
            for di, kun in enumerate(KUNLAR):
                sc = di * PERIODS_PER_DAY + 2
                ws.merge_cells(start_row=3, start_column=sc, end_row=3, end_column=sc + PERIODS_PER_DAY - 1)
                c = ws.cell(3, sc, kun); c.font = hfont; c.fill = hf
                c.alignment = Alignment(horizontal='center', vertical='center')
            for di in range(6):
                for pp in range(PERIODS_PER_DAY):
                    c = ws.cell(4, di * PERIODS_PER_DAY + pp + 2, str(pp + 1))
                    c.font = pfont; c.fill = pf; c.alignment = Alignment(horizontal='center')

            for ci, cls in enumerate(self.classes):
                rn = ci + 5
                nc = ws.cell(rn, 1, cls[1])
                nc.font = Font(bold=True, size=9, color="FFFFFF")
                nc.fill = PatternFill(start_color="34495E", end_color="34495E", fill_type="solid")
                nc.alignment = Alignment(horizontal='center', vertical='center')
                for day in range(6):
                    for period in range(PERIODS_PER_DAY):
                        info = tt_data.get((cls[0], day, period), {})
                        val = info.get('subject_name', '') if info else ''
                        c = ws.cell(rn, day * PERIODS_PER_DAY + period + 2, val)
                        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                        c.font = Font(size=8); c.border = tb

            ws.column_dimensions['A'].width = 10
            for col in range(2, 6 * PERIODS_PER_DAY + 2):
                ws.column_dimensions[get_column_letter(col)].width = 12
        wb.save(fn)

    def _save_word(self, fn):
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT

        tt_list = self._get_export_tt()
        doc = Document()

        et, eid, ename = self._get_type()
        school = self.school.text()

        week_names = ["1-hafta (Toq)", "2-hafta (Juft)"]
        for wi, tt_data in enumerate(tt_list):
            ws_name = week_names[wi] if wi < len(week_names) else f"Hafta {wi+1}"

            # Sahifa ajratish (ikkala hafta uchun)
            if wi > 0:
                doc.add_page_break()

            # Sarlavha
            title = doc.add_heading(f'DARS JADVALI — {ws_name}', level=0)
            title.alignment = WD_TABLE_ALIGNMENT.CENTER

            if school:
                p = doc.add_paragraph(school)
                p.alignment = WD_TABLE_ALIGNMENT.CENTER
                p.style.font.size = Pt(12)

            doc.add_paragraph()

            if et == 'umumiy':
                self._add_word_grid(doc, tt_data, self.classes)
            elif et == 'sinf':
                cls = next((c for c in self.classes if c[0] == eid), None)
                if cls:
                    doc.add_heading(f"Sinf: {cls[1]}", level=1)
                    self._add_word_single_class(doc, tt_data, cls)
            elif et == 'ustoz':
                doc.add_heading(f"O'qituvchi: {ename}", level=1)
                self._add_word_teacher(doc, tt_data, eid, ename)
            elif et == 'fan':
                doc.add_heading(f"Fan: {ename}", level=1)
                self._add_word_subject(doc, tt_data, ename)

        doc.save(fn)

    def _add_word_grid(self, doc, tt, classes):
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT

        for cls in classes:
            doc.add_heading(f"Sinf: {cls[1]}", level=2)
            self._add_word_single_class(doc, tt, cls)
            doc.add_paragraph()

    def _add_word_single_class(self, doc, tt, cls):
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT

        table = doc.add_table(rows=PERIODS_PER_DAY + 1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Sarlavha
        headers = ['', 'Dushanba', 'Seshanba', 'Chorshanba', 'Payshanba', 'Juma', 'Shanba']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        # Darslar
        for period in range(PERIODS_PER_DAY):
            table.cell(period + 1, 0).text = str(period + 1)
            for day in range(6):
                info = tt.get((cls[0], day, period), {})
                val = info.get('subject_name', '') if info else ''
                table.cell(period + 1, day + 1).text = val

    def _add_word_teacher(self, doc, tt, teacher_id, teacher_name):
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT

        table = doc.add_table(rows=PERIODS_PER_DAY + 1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['', 'Dushanba', 'Seshanba', 'Chorshanba', 'Payshanba', 'Juma', 'Shanba']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        for period in range(PERIODS_PER_DAY):
            table.cell(period + 1, 0).text = str(period + 1)
            for day in range(6):
                for key, info in tt.items():
                    if info.get('teacher_id') == teacher_id and key[1] == day and key[2] == period:
                        table.cell(period + 1, day + 1).text = info.get('class_name', '')
                        break

    def _add_word_subject(self, doc, tt, subject_name):
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT

        table = doc.add_table(rows=PERIODS_PER_DAY + 1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['', 'Dushanba', 'Seshanba', 'Chorshanba', 'Payshanba', 'Juma', 'Shanba']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        for period in range(PERIODS_PER_DAY):
            table.cell(period + 1, 0).text = str(period + 1)
            for day in range(6):
                for key, info in tt.items():
                    if info.get('subject_name') == subject_name and key[1] == day and key[2] == period:
                        table.cell(period + 1, day + 1).text = info.get('class_name', '')
                        break

    def _save_html(self, fn):
        html = self._build_html()
        with open(fn, 'w', encoding='utf-8') as f:
            f.write(html)


